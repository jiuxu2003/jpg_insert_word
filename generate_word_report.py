"""
Generate a Word document that lays out images in two columns with numbered captions.

Usage:
    python generate_word_report.py --images ./pics --output report.docx --per-row 2

The caption text defaults to "图{idx} {stem}", where stem is taken from each image's file name.
"""

from __future__ import annotations

import argparse
import io
from pathlib import Path
import re
from dataclasses import dataclass
from typing import Callable, Iterable, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tif", ".tiff", ".webp"}


@dataclass
class ImageItem:
    path: Path
    mileage_text: str
    mileage_value: float
    y_value: int


def parse_mileage(text: str) -> float:
    digits = re.findall(r"[0-9]+(?:\.[0-9]+)?", text)
    if not digits:
        return float("inf")
    try:
        return float(digits[0])
    except ValueError:
        return float("inf")


def find_images(folder: Path) -> List[ImageItem]:
    items: List[ImageItem] = []
    for path in folder.iterdir():
        if path.is_file() and path.suffix.lower() in IMAGE_EXTS:
            stem = path.stem
            if "-" in stem:
                parts = stem.split("-", 1)
                mileage_text = parts[0].strip()
                try:
                    y_value = int(parts[1].strip())
                except (ValueError, IndexError):
                    y_value = 0
            else:
                mileage_text = stem.strip()
                y_value = 0
            items.append(
                ImageItem(
                    path=path,
                    mileage_text=mileage_text,
                    mileage_value=parse_mileage(mileage_text),
                    y_value=y_value,
                )
            )
    items.sort(key=lambda item: (item.mileage_value, item.y_value, item.mileage_text))
    return items


def chunked(items: Iterable[ImageItem], size: int) -> Iterable[List[ImageItem]]:
    row: List[ImageItem] = []
    for item in items:
        row.append(item)
        if len(row) == size:
            yield row
            row = []
    if row:
        yield row


def add_caption_run(run, label: str) -> None:
    run.text = label
    font = run.font
    font.size = Pt(10.5)
    font.name = "Times New Roman"
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def _cm_to_twips(value: float) -> int:
    return int(round(value * 567))


def set_cell_margins(cell, *, top=0.0, start=0.0, bottom=0.0, end=0.0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for attr, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tcMar.find(qn(f"w:{attr}"))
        if element is None:
            element = OxmlElement(f"w:{attr}")
            tcMar.append(element)
        element.set(qn("w:type"), "dxa")
        element.set(qn("w:w"), str(_cm_to_twips(value)))


def prepare_image_stream(image_path: Path, width_cm: float, height_cm: float) -> io.BytesIO:
    dpi = 220
    target_w = int(round(width_cm / 2.54 * dpi))
    target_h = int(round(height_cm / 2.54 * dpi))

    with Image.open(image_path) as img:
        img = img.convert("RGB")
        scale = min(target_w / img.width, target_h / img.height, 1.0)
        new_size = (
            max(1, int(round(img.width * scale))),
            max(1, int(round(img.height * scale))),
        )
        resized = img.resize(new_size, Image.LANCZOS)

        canvas = Image.new("RGB", (target_w, target_h), color="white")
        offset = ((target_w - new_size[0]) // 2, (target_h - new_size[1]) // 2)
        canvas.paste(resized, offset)

        stream = io.BytesIO()
        canvas.save(stream, format="PNG", dpi=(dpi, dpi))
        stream.seek(0)
        return stream


def add_image_block(cell, image_stream: io.BytesIO, width_cm: float, height_cm: float) -> None:
    cell.text = ""
    pic_para = cell.paragraphs[0]
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.space_before = Pt(0)
    pic_para.paragraph_format.space_after = Pt(2)
    run = pic_para.add_run()
    run.add_picture(image_stream, width=Cm(width_cm), height=Cm(height_cm))


def build_document(
    images: List[ImageItem],
    per_row: int,
    width_cm: float,
    height_cm: float,
    progress_callback: Optional[Callable[[int, int], None]] = None,
) -> Document:
    document = Document()
    section = document.sections[0]
    section.left_margin = Cm(2.76)
    section.right_margin = Cm(2.76)

    table = document.add_table(rows=0, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    gap_width = max(0.1, 11.70 - 2.76 - width_cm)
    column_widths = [width_cm, gap_width, width_cm]

    def add_row():
        row = table.add_row()
        for idx, width in enumerate(column_widths):
            cell = row.cells[idx]
            cell.width = Cm(width)
            set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        return row

    image_index = 1
    processed = 0
    total = len(images)
    image_streams: List[io.BytesIO] = []
    for row_images in chunked(images, per_row):
        img_row = add_row()
        cap_row = add_row()

        for col_idx, cell_idx in enumerate((0, 2)):
            if col_idx < len(row_images):
                image_item = row_images[col_idx]
                s_cycle = (image_index - 1) % 3 + 1
                label = (
                    f"图5.6-{image_index} {image_item.mileage_text}S{s_cycle}沉降曲线"
                )
                stream = prepare_image_stream(image_item.path, width_cm, height_cm)
                image_streams.append(stream)
                add_image_block(img_row.cells[cell_idx], stream, width_cm, height_cm)
                processed += 1
                if progress_callback:
                    progress_callback(processed, total)

                caption_cell = cap_row.cells[cell_idx]
                caption_cell.text = ""
                caption_para = caption_cell.paragraphs[0]
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.paragraph_format.space_before = Pt(0)
                caption_para.paragraph_format.space_after = Pt(0)
                add_caption_run(caption_para.add_run(), label)
                image_index += 1
            else:
                img_row.cells[cell_idx].text = ""
                cap_row.cells[cell_idx].text = ""

        img_row.cells[1].text = ""
        cap_row.cells[1].text = ""

    # Remove table borders
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "nil")

    return document


def generate_word_report(
    image_dir: Path,
    output_path: Path,
    per_row: int = 2,
    width_cm: float = 7.6,
    height_cm: float = 4.7,
    progress_callback: Optional[Callable[[int, int], None]] = None,
) -> None:
    if not image_dir.is_dir():
        raise FileNotFoundError(f"Image directory not found: {image_dir}")

    images = find_images(image_dir)
    if not images:
        raise ValueError("No supported image files were found in the folder.")

    document = build_document(
        images, per_row, width_cm, height_cm, progress_callback=progress_callback
    )
    document.save(str(output_path))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a Word doc from images.")
    parser.add_argument("--images", required=True, type=Path, help="Folder with images")
    parser.add_argument(
        "--output",
        default="图片汇总.docx",
        type=Path,
        help="Output Word file name",
    )
    parser.add_argument(
        "--per-row",
        type=int,
        default=2,
        help="How many images to place per row",
    )
    parser.add_argument("--width-cm", type=float, default=7.6, help="Image width in cm")
    parser.add_argument("--height-cm", type=float, default=4.7, help="Image height in cm")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    progress_cb: Optional[Callable[[int, int], None]] = None

    try:
        import sys

        if sys.stdout.isatty():
            def cli_progress(done: int, total: int) -> None:
                percent = done / total * 100 if total else 0
                print(f"\r正在处理图片：{done}/{total} ({percent:5.1f}%)", end="", flush=True)

            progress_cb = cli_progress
    except Exception:  # noqa: BLE001
        progress_cb = None

    try:
        generate_word_report(
            args.images,
            args.output,
            args.per_row,
            args.width_cm,
            args.height_cm,
            progress_callback=progress_cb,
        )
    except Exception as exc:
        raise SystemExit(str(exc))
    else:
        if progress_cb is not None:
            print()
        print("Saved images to", args.output)


if __name__ == "__main__":
    main()

